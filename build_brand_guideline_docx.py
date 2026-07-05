"""Build the Digital Brand Guideline as a DOCX companion to the HTML."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Palette
MAROON = RGBColor(0x79, 0x21, 0x1C)
MAROON_INK = RGBColor(0x2A, 0x0E, 0x0C)
RUST = RGBColor(0x96, 0x29, 0x01)
GOLD = RGBColor(0xC8, 0x9B, 0x3C)
GRAPHITE = RGBColor(0x33, 0x33, 0x33)
INK = RGBColor(0x1B, 0x1B, 0x1A)
MID = RGBColor(0x7A, 0x73, 0x68)
AGRI = RGBColor(0x1B, 0x7A, 0x3D)
PAPER2 = "F6F6F4"
LINE = "D9D2C4"

doc = Document()

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)


def shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_para(text, size=11, color=INK, bold=False, italic=False, space_after=6, align=None, font="Calibri"):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading(text, level=1):
    if level == 0:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Cambria"
        r.font.size = Pt(28)
        r.font.color.rgb = MAROON_INK
        r.bold = False
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(6)
    elif level == 1:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Cambria"
        r.font.size = Pt(20)
        r.font.color.rgb = MAROON_INK
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(13)
        r.font.color.rgb = MAROON
        r.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_kicker(text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = RUST
    r.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_bullets(items, size=11, color=GRAPHITE):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.font.color.rgb = color


def add_do_dont(do_items, dont_items):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = True
    hdr = tbl.rows[0].cells
    hdr[0].text = ""
    hdr[1].text = ""

    # Do cell
    c0 = hdr[0]
    shade(c0, "F4F9F4")
    p = c0.paragraphs[0]
    r = p.add_run("DO")
    r.font.name = "Calibri"; r.font.size = Pt(9); r.font.color.rgb = AGRI; r.bold = True
    for it in do_items:
        p2 = c0.add_paragraph()
        p2.paragraph_format.left_indent = Pt(0)
        rr = p2.add_run("•  " + it)
        rr.font.name = "Calibri"; rr.font.size = Pt(10); rr.font.color.rgb = GRAPHITE

    # Don't cell
    c1 = hdr[1]
    shade(c1, "FCF3EF")
    p = c1.paragraphs[0]
    r = p.add_run("DON’T")
    r.font.name = "Calibri"; r.font.size = Pt(9); r.font.color.rgb = RUST; r.bold = True
    for it in dont_items:
        p2 = c1.add_paragraph()
        rr = p2.add_run("•  " + it)
        rr.font.name = "Calibri"; rr.font.size = Pt(10); rr.font.color.rgb = GRAPHITE

    add_para("", size=6, space_after=6)


def add_swatch_row(rows):
    """rows: list of (name, hex, role)"""
    tbl = doc.add_table(rows=len(rows) + 1, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["", "Token", "Hex", "Role"]):
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Calibri"; r.font.size = Pt(9); r.bold = True; r.font.color.rgb = MID
    for i, (name, hexv, role) in enumerate(rows, start=1):
        cells = tbl.rows[i].cells
        shade(cells[0], hexv.lstrip("#"))
        for j, val in enumerate([name, hexv, role], start=1):
            p = cells[j].paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = INK
            if j == 2:
                r.font.name = "Consolas"; r.font.color.rgb = MAROON
    add_para("", size=6, space_after=6)


# ==================== COVER ====================
add_kicker("Digital Brand Guideline · Version 1")
add_heading("Rules for the digital face of Pantnagar.", level=0)
add_para(
    "A working document for the team maintaining and extending gbpuat.ac.in. It records the design system already built on the modernised website — colour, type, layout, components, motion, accessibility, bilingual rules — and the reasons behind each choice. Read it before adding a page, changing a component, or bringing on a new contributor.",
    size=12, color=GRAPHITE, italic=False, space_after=14, font="Cambria"
)

# Cover meta table
meta_tbl = doc.add_table(rows=2, cols=2)
meta = [
    ("Institution", "G.B. Pant University of Agriculture & Technology"),
    ("Scope", "Website only — gbpuat.ac.in"),
    ("Companion to", "Institutional Brand Guideline (AS-IS)"),
    ("Status", "Working draft for handover"),
]
for i in range(2):
    for j in range(2):
        label, value = meta[i * 2 + j]
        cell = meta_tbl.rows[i].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(label.upper())
        r.font.name = "Calibri"; r.font.size = Pt(8); r.font.color.rgb = GOLD; r.bold = True
        p2 = cell.add_paragraph()
        r2 = p2.add_run(value)
        r2.font.name = "Cambria"; r2.font.size = Pt(11); r2.font.color.rgb = INK

doc.add_page_break()

# ==================== CONTENTS ====================
add_heading("Contents", level=1)
toc = [
    "1. Foreword — what this document is",
    "2. What changed, and why",
    "3. Design principles",
    "4. Colour system",
    "5. Typography",
    "6. Logo and crest",
    "7. Layout and rhythm",
    "8. Components",
    "9. Motion",
    "10. Accessibility",
    "11. Bilingual rules",
    "12. Content voice",
    "13. Governance",
]
for entry in toc:
    add_para(entry, size=11, color=INK, space_after=4)
doc.add_page_break()

# ==================== 1. FOREWORD ====================
add_kicker("Chapter One")
add_heading("What this document is, and is not.", level=0)
add_para(
    "The AS-IS Institutional Brand Guideline governs the identity of the university — its crest, its motto, its formal colour and typography. This document sits alongside it and governs one thing only: how that identity is expressed on the website.",
    size=12, color=GRAPHITE, font="Cambria"
)
add_heading("What this document governs", level=2)
add_bullets([
    "The design tokens actually used on gbpuat.ac.in — colours, fonts, spacing, rules, shadows.",
    "The components already built and shipped — headers, cards, ticker, notices, footer, gallery.",
    "The rules a new contributor must follow when adding a page or a section.",
    "The anti-patterns that must not return — Bootstrap blue, autoplay carousels, right-click blocking, generic AI-template aesthetics.",
])
add_heading("What this document does not govern", level=2)
add_bullets([
    "Print, stationery, or signage — refer to the AS-IS Institutional Brand Guideline.",
    "The crest as a mark of the institution — the AS-IS guideline is authoritative on clear space, minimum size, and formal usage.",
    "Social media identity, email templates, or event branding.",
])
add_heading("How to read it", level=2)
add_para(
    "Each chapter starts with the principle, then gives the specific tokens and rules that follow from it, then closes with a Do and Don’t for the team. When this document contradicts the AS-IS guideline on a website-specific question, this document applies. On any question of institutional identity, the AS-IS guideline applies.",
    size=11, color=GRAPHITE,
)

# ==================== 2. DELTA ====================
add_kicker("Chapter Two")
add_heading("What changed on the new site — and the reason.", level=0)
add_para(
    "The redesign is not a repaint. It replaces a set of default choices that had accumulated over the years with a system that reads as bespoke to Pantnagar. The table below records the substantive shifts, so the team knows which anti-patterns to reject when they resurface.",
    size=12, color=GRAPHITE, font="Cambria"
)

delta_rows = [
    ("Primary accent", "Bootstrap default blue (#0D6EFD)", "Institutional Maroon (#79211C)", "The maroon is Pantnagar’s own. The default blue is nobody’s."),
    ("Body typeface", "Arial / system sans defaults", "IBM Plex Sans + EB Garamond", "The site now reads as a serious research university, not a generic template."),
    ("Hindi type", "Same font as English, poor render", "Noto Serif / Sans Devanagari, at parity", "Hindi must not appear typographically demoted next to English."),
    ("Hero", "Autoplaying jQuery slider", "Editorial hero + measured carousel with progress bar", "Autoplay steals attention and fails accessibility."),
    ("Right-click / copy", "Blocked", "Allowed", "A public institution’s content must be quotable, printable, screen-readable."),
    ("Legacy browser notice", "Shown", "Removed", "Modern evergreen browsers are the baseline. The notice was noise."),
    ("Palette range", "Corporate primaries", "Earth tones — maroon, rust, gold, agri green", "Grounded in the agricultural and heritage identity of the institution."),
    ("Background", "Flat white", "Warm paper with a faint grain", "Editorial atmosphere. The site should feel printed, not projected."),
    ("Security headers", "Absent or minimal", "CSP, referrer policy, X-Content-Type-Options", "Defence in depth for a public institution’s public face."),
]

tbl = doc.add_table(rows=len(delta_rows) + 1, cols=4)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, h in enumerate(["Element", "Old site", "New site", "Reason"]):
    p = hdr[i].paragraphs[0]
    r = p.add_run(h)
    r.font.name = "Calibri"; r.font.size = Pt(9); r.bold = True; r.font.color.rgb = MID
for i, row in enumerate(delta_rows, start=1):
    cells = tbl.rows[i].cells
    for j, val in enumerate(row):
        p = cells[j].paragraphs[0]
        r = p.add_run(val)
        r.font.name = "Calibri"; r.font.size = Pt(10)
        if j == 1:
            r.font.color.rgb = RUST
        elif j == 2:
            r.font.color.rgb = AGRI; r.bold = True
        else:
            r.font.color.rgb = INK
add_para("", size=6, space_after=6)
add_para(
    "If a request lands to bring any of the old defaults back, treat it as a policy question, not a design one. The reason column is the answer.",
    size=11, color=MAROON_INK, italic=True, font="Cambria"
)

# ==================== 3. PRINCIPLES ====================
add_kicker("Chapter Three")
add_heading("Six principles that decide the close calls.", level=0)
add_para(
    "When a specification is silent, these principles decide. They are ordered — earlier principles outrank later ones.",
    size=12, color=GRAPHITE, font="Cambria"
)
principles = [
    ("1. Editorial, not corporate.", "The site should read as a serious institutional publication — a museum publication, a heritage magazine, a research monograph. It should not read as a SaaS marketing page. When choosing between a designed page and a designed template, choose the designed page."),
    ("2. Heritage, not trend.", "Pantnagar was founded in 1960. The design language should be legible fifteen years from now. Reject any device whose main appeal is being current."),
    ("3. Paper first, screen second.", "The neutral base is warm paper, not clinical white. The background carries a faint grain. Rules are hair-thin. Whitespace is generous. If a section feels crowded on screen, it will feel unreadable in print, and it must be pared back."),
    ("4. Devanagari at parity.", "Where Hindi appears, it appears in a matched serif or sans — Noto Serif Devanagari or IBM Plex Sans Devanagari — at a size and weight that reads as equal to the Latin. Hindi is never smaller, thinner, or lower in the visual hierarchy than English."),
    ("5. Restraint in motion.", "Motion serves reading. A hover translate, an opacity crossfade, a soft image zoom over eight seconds — these are the permitted vocabulary. Autoplay, parallax scroll, marquee tickers, and decorative animation on entry are not."),
    ("6. No AI-template defaults.", "The site must never read as machine-generated. No Inter, Roboto, Arial, or system-ui as display type; no purple-indigo gradients on white; no three-card feature row with soft pastel icons; no cookie-cutter hero with a large button in the middle. When a proposed component looks like every other AI-built page on the internet, it is wrong for this site."),
]
for h, body in principles:
    add_heading(h, level=2)
    add_para(body, size=11, color=GRAPHITE)

# ==================== 4. COLOUR ====================
add_kicker("Chapter Four")
add_heading("The palette, and where each colour is allowed.", level=0)
add_para(
    "Every colour on the site is a token defined in the stylesheet. New colours are not added in a page; they are added to the tokens. The palette is intentionally narrow.",
    size=12, color=GRAPHITE, font="Cambria"
)

add_heading("Institutional core", level=2)
add_swatch_row([
    ("Institutional Maroon", "#79211C", "Primary. CTAs, brand text, primary navigation, section rules."),
    ("Maroon Deep", "#5A1714", "Hover state for maroon surfaces."),
    ("Maroon Ink", "#2A0E0C", "Dark section backgrounds, primary heading colour on paper."),
    ("Earthen Rust", "#962901", "Kickers, dates, secondary accents. Never for body text or large fills."),
])

add_heading("Accents", level=2)
add_swatch_row([
    ("Wheat Gold", "#C89B3C", "Thin rules under kickers, active-nav underline, numeric marks."),
    ("Gold Soft", "#E8CB86", "Text on maroon-ink surfaces; italic labels in dark sections."),
    ("Agri Green", "#1B7A3D", "Positive states, agriculture-context tags. Rare use."),
    ("Link Steel", "#2C67A0", "Reserved for a generic link where maroon would be wrong. Almost never needed."),
])

add_heading("Neutrals — warm, paper-first", level=2)
add_swatch_row([
    ("Paper", "#FFFFFF", "Base card and section background."),
    ("Paper 2", "#F6F6F4", "Alternating section background — the warm off-white that gives the paper feel."),
    ("Wash", "#ECECEC", "Muted fills where a neutral wash is needed."),
    ("Line", "#D9D2C4", "Standard 1px rules and card borders."),
    ("Line Soft", "#E7E1D3", "Softer rules where 1px on Line would be too strong."),
    ("Mid", "#7A7368", "Uppercase labels, kicker meta, muted secondary text."),
    ("Graphite", "#333333", "Body copy on paper."),
    ("Ink", "#1B1B1A", "Highest-emphasis body text."),
])

add_heading("Contrast floors", level=2)
add_para(
    "Every text-and-background pairing must meet WCAG 2.1 AA — 4.5:1 for body text, 3:1 for large text (18pt+ or 14pt bold+). The pairings below have been verified; any new pairing must be verified before ship.",
    size=11, color=GRAPHITE
)
contrast_rows = [
    ("Ink #1B1B1A", "Paper #FFFFFF", "18.9 : 1", "Body text on paper."),
    ("Graphite #333333", "Paper #FFFFFF", "12.6 : 1", "Secondary body on paper."),
    ("Maroon #79211C", "Paper #FFFFFF", "9.0 : 1", "Links, CTAs, headings."),
    ("Mid #7A7368", "Paper #FFFFFF", "4.6 : 1", "Labels — meets body floor."),
    ("Paper #FFFFFF", "Maroon #79211C", "9.0 : 1", "CTA button text."),
    ("Gold Soft #E8CB86", "Maroon Ink #2A0E0C", "10.7 : 1", "Kickers on dark sections."),
]
tbl = doc.add_table(rows=len(contrast_rows) + 1, cols=4)
tbl.style = "Light Grid Accent 1"
for i, h in enumerate(["Foreground", "Background", "Ratio", "Use"]):
    p = tbl.rows[0].cells[i].paragraphs[0]
    r = p.add_run(h); r.font.name = "Calibri"; r.font.size = Pt(9); r.bold = True; r.font.color.rgb = MID
for i, row in enumerate(contrast_rows, start=1):
    for j, val in enumerate(row):
        p = tbl.rows[i].cells[j].paragraphs[0]
        r = p.add_run(val); r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = INK
add_para("", size=6, space_after=6)

add_do_dont(
    [
        "Use maroon for anything the user must identify as institutional — CTAs, active nav, section rules.",
        "Use rust for kickers and dates — the small uppercase labels that anchor an editorial layout.",
        "Use gold as a line, not a fill. It is the finishing detail of the system.",
        "Alternate paper and paper-2 to give the page rhythm across sections.",
    ],
    [
        "Never use Bootstrap blue (#0D6EFD) or any cold corporate primary.",
        "Never introduce a purple, indigo, or teal gradient on white — that is the AI-template signature.",
        "Never fill a large area with gold. It reads as ornament.",
        "Never use rust for body text — it is a label colour, not a reading colour.",
        "Never introduce a new brand colour without adding it to the tokens.",
    ],
)

# ==================== 5. TYPOGRAPHY ====================
add_kicker("Chapter Five")
add_heading("Four families, and the discipline of how they meet.", level=0)
add_para(
    "The site pairs a heritage academic serif with a humanist grotesk, and mirrors both in Devanagari. No fifth family is introduced. Weights are limited. Italic is used with intent.",
    size=12, color=GRAPHITE, font="Cambria"
)

add_heading("The families", level=2)
add_bullets([
    "EB Garamond — display and long-form. Weights 400 and 500; italic for emphasis. The academic voice of the site.",
    "IBM Plex Sans — UI, body, labels, navigation. Weights 300 to 600. The functional voice of the site.",
    "Noto Serif Devanagari — Hindi display and headings. Matches Garamond in role.",
    "IBM Plex Sans Devanagari — Hindi body and UI. Matches Plex Sans in role.",
])
add_para(
    "All four are loaded from Google Fonts with display=swap. If Google Fonts must be self-hosted for policy reasons, the same four families in the same weights are the only substitution permitted.",
    size=11, color=GRAPHITE
)

add_heading("The type scale actually used", level=2)
add_para(
    "All display sizes use clamp() so type breathes between phones and desktops. The values below are the shipped tokens — do not hand-tune sizes in a component.",
    size=11, color=GRAPHITE
)
scale_rows = [
    ("Hero H1", "Serif", "clamp(44px, 6.5vw, 82px)", "300", "0.98"),
    ("Section H2", "Serif", "clamp(30px, 3.4vw, 44px)", "350", "1.08"),
    ("Card H3", "Serif", "22 to 28px", "400", "1.10 to 1.20"),
    ("Lede", "Serif", "clamp(18px, 1.6vw, 22px)", "350 to 400", "1.45"),
    ("Body", "Sans", "16.5 to 17px", "400", "1.55"),
    ("Meta / label", "Sans", "11 to 13px", "500 to 600", "1.35"),
    ("Kicker", "Sans, uppercase", "11 to 12px", "500 to 600", "1"),
]
tbl = doc.add_table(rows=len(scale_rows) + 1, cols=5)
tbl.style = "Light Grid Accent 1"
for i, h in enumerate(["Role", "Family", "Size", "Weight", "Line-height"]):
    p = tbl.rows[0].cells[i].paragraphs[0]
    r = p.add_run(h); r.font.name = "Calibri"; r.font.size = Pt(9); r.bold = True; r.font.color.rgb = MID
for i, row in enumerate(scale_rows, start=1):
    for j, val in enumerate(row):
        p = tbl.rows[i].cells[j].paragraphs[0]
        r = p.add_run(val); r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = INK
        if j == 2:
            r.font.name = "Consolas"; r.font.color.rgb = MAROON
add_para("", size=6, space_after=6)

add_heading("Discipline", level=2)
add_bullets([
    "Headings are set at weight 300 to 400 — the serif carries authority through its form, not its weight.",
    "Letter-spacing on display is negative (-0.015em to -0.028em).",
    "Kickers and labels are uppercase, tracked (letter-spacing 0.18em to 0.28em).",
    "Italic is used in the serif only, for editorial emphasis — never in the sans.",
    "Old-style numerals on display statistics; tabular numerals on tables and dates.",
])

add_do_dont(
    [
        "Set every heading in EB Garamond (or Noto Serif Devanagari for Hindi).",
        "Set every kicker as uppercase Plex Sans with 1px gold underline.",
        "Use [lang='hi'] to switch families automatically — never hard-code Hindi styles inline.",
        "Trust clamp(). Do not add a new fixed-size heading class.",
    ],
    [
        "Never introduce Inter, Roboto, Arial, system-ui, Space Grotesk, Lato, or Poppins — the AI-template tell.",
        "Never set a heading in weight 700+ — the serif at heavy weights collapses the editorial feel.",
        "Never set Hindi in the Latin serif or sans — the render is wrong.",
        "Never underline body text for emphasis — use italic or the maroon inline.",
    ],
)

# ==================== 6. LOGO ====================
add_kicker("Chapter Six")
add_heading("The crest on screen.", level=0)
add_para(
    "The AS-IS Institutional Brand Guideline is authoritative on the crest as a mark. This chapter records only how it is used on the website.",
    size=12, color=GRAPHITE, font="Cambria"
)
add_heading("Sizes", level=2)
add_bullets([
    "Masthead crest — 72 by 72 CSS pixels. Rendered from the optimised PNG (logo_small.png).",
    "Hero crest plate — 190 CSS pixels wide, centred, with a soft drop shadow.",
    "Favicon — the same PNG, served as rel='icon'.",
])
add_heading("Backgrounds", level=2)
add_bullets([
    "The crest appears on paper (#FFFFFF or #F6F6F4). It is not used on maroon-ink backgrounds — those sections use the wordmark or the italic label instead.",
    "The crest is never placed on a photograph. The hero uses a crest plate on paper, not on the image.",
])
add_heading("Format", level=2)
add_para(
    "The current asset is a PNG. Convert to SVG at the first opportunity — the crest is used at multiple sizes and the SVG will render crisply on high-DPI screens without shipping multiple bitmaps. The SVG must be reviewed against the master before it replaces the PNG.",
    size=11, color=GRAPHITE
)
add_heading("Cross-origin images (referrer policy)", level=2)
add_para(
    "When the site pulls an image from a third-party host (for example, an officer photograph from a government page), add referrerpolicy='no-referrer' to the img tag. Some hosts block hot-linking when a referrer is sent; the no-referrer policy is what keeps those images resolving in Safari and iOS.",
    size=11, color=GRAPHITE
)
add_do_dont(
    [
        "Use the official crest asset, unmodified, from the repository.",
        "Give the crest at least its own height as clear space on all sides.",
        "Refer to the AS-IS guideline for anything beyond web usage.",
    ],
    [
        "Never recolour, outline, or stylise the crest.",
        "Never place the crest on a photograph or a coloured field.",
        "Never render the crest below 40 CSS pixels — it becomes illegible.",
    ],
)

# ==================== 7. LAYOUT ====================
add_kicker("Chapter Seven")
add_heading("The grid, the rhythm, and the little details that carry it.", level=0)
add_para(
    "Everything on the site sits inside one measured container. Sections alternate paper and paper-2. Rules are hair-thin. The maroon is used at 3px only where it must anchor.",
    size=12, color=GRAPHITE, font="Cambria"
)
add_heading("The wrapper", level=2)
add_bullets([
    "Max width 1200px, centred.",
    "Horizontal padding clamp(1rem, 2.5vw, 2rem).",
    "Section vertical padding clamp(60px, 8vw, 96px) top and bottom.",
])
add_heading("The spacing scale", level=2)
add_para(
    "Multiples of 4px are used throughout. The scale in practice: 4, 8, 12, 16, 20, 22, 24, 28, 32, 40, 48, 56, 64, 80, 96. Do not invent intermediate values.",
    size=11, color=GRAPHITE
)
add_heading("Rules and borders", level=2)
add_bullets([
    "Line (#D9D2C4) — 1px. Standard rule.",
    "Line Soft (#E7E1D3) — 1px. Softer rule where standard is too heavy.",
    "Maroon 3px — reserved for anchor rules under the primary nav, above the notices strip, and at the start of a colleges directory.",
    "Gold 1px — under kickers, under the active-nav underline, under dot indicators.",
])
add_heading("Shadows", level=2)
add_para(
    "Shadows are used sparingly and are always warm. The permitted shadow is drop-shadow(0 6px 16px rgba(42, 14, 12, 0.12)) — a soft maroon-tinted glow that reads as printed relief. Cool grey shadows are forbidden.",
    size=11, color=GRAPHITE
)
add_heading("Paper grain", level=2)
add_para(
    "A double-layer radial-gradient noise sits behind the entire page at 50% opacity, using mix-blend-mode: multiply. It is present but never announced. If a section feels flat, do not add gradients — check that the grain is rendering.",
    size=11, color=GRAPHITE
)
add_heading("Borders and radii", level=2)
add_bullets([
    "Buttons: border-radius 2px. The site uses very slight rounding — never pill buttons, never fully square.",
    "Search chip: border-radius 999px. The only pill in the system.",
    "Cards: no radius. The editorial feel comes from the sharp corners plus the rule border.",
])

# ==================== 8. COMPONENTS ====================
add_kicker("Chapter Eight")
add_heading("The component library, one page at a time.", level=0)
add_para(
    "The site is built from a fixed set of components. New pages are assembled from this set. A new component is added only when nothing in the set can be reasonably adapted.",
    size=12, color=GRAPHITE, font="Cambria"
)

components = [
    ("Utility bar", "The topmost strip. Maroon-ink background, gold-soft text. Contains accessibility controls (A+ / A-), quick links (Web Mail, Weather Advisory, Alumni, Success Stories), and the language toggle. Height 36px minimum, font-size 13px."),
    ("Masthead", "Paper background, 1px line-bottom. The 72px crest sits left, bilingual wordmark centred against it, search chip and Admissions CTA on the right. On mobile the CTA collapses first; the search chip second."),
    ("Primary navigation", "Sits under the masthead. 3px maroon bottom border anchors it. Active link is marked in maroon with a gold 3px underline. Hover moves to maroon."),
    ("Hero carousel", "Full-bleed image with a warm dark gradient overlay. Bottom-anchored caption with a gold kicker and a serif title. Filmstrip thumbnails below the frame. Progress bar in gold at 2px. Reduced motion is respected — the slow zoom and progress bar are removed for users who request it."),
    ("Hero (non-carousel)", "Asymmetric grid: title-and-lede on the left, crest plate with motto and metadata on the right, separated by a 1px vertical rule. Title at clamp(44px, 6.5vw, 82px), weight 300, letter-spacing -0.028em. The rust italic emphasis inside the title is the site’s editorial signature."),
    ("Notices strip", "Maroon-ink field. Italic serif label ‘Notices’ in gold-soft, a gold count badge next to it, five-column notice list, and a ‘See all’ link. Bottom bordered in 3px gold. This replaces the old scrolling ticker — no marquee."),
    ("Section heading", "Every section opens with a rust kicker, a serif H2 in maroon-ink weight 350, and a short graphite sub-line. A ‘See more’ link may sit top-right of the heading block."),
    ("Statistic display", "An oversized old-style numeral in the serif at clamp(130px, 18vw, 260px), paired with a short caption and a small dl of supporting figures on the right. The numeral carries the emphasis — the label is quiet."),
    ("Card — quick access", "Paper background, 1px line-soft border, 3px maroon left border. Numbered kicker in gold at the top. Serif title, sans description, uppercase CTA line at the bottom over a hairline top rule. Hover translates 3px up and shifts the left border to gold."),
    ("College card / directory row", "The colleges are listed as a numbered editorial directory — 22px serif number in gold on the left, 22px serif title in maroon-ink, one-line graphite description. 1px line-soft between rows, opened by a 2px maroon rule above. Hover slides the row 4px right."),
    ("News — editorial layout", "One lead story on the left (16:10 image, kicker, date, serif headline, lede) and a numbered rail on the right (date, serif one-liner, hairline rules between items). No card grid. This is the site’s distinctive news treatment — it must be preserved."),
    ("Gallery card", "Placed in the maroon-ink gallery section. Semi-transparent white fill, 1px semi-transparent border, gold kicker, white serif H3, muted body. Hover lifts the fill to 6% white and shifts the border to gold."),
    ("Trust / important links row", "Grid of five equal cells with hairline dividers. On mobile collapses to three then two columns while keeping the internal divider grid tidy."),
    ("Circulars / notices list", "Two-column row: date on the left (140px, rust, uppercase, tabular), title on the right. Opened by 2px maroon rule. Hover shifts the title to maroon."),
    ("Footer", "Maroon-ink field, gold 3px top border. Institutional wordmark, address block, contact block, and a compact link column. Reserves gold-soft for section labels and paper-white for the wordmark."),
    ("Buttons", "Primary CTA — background maroon, colour white, padding 11px 18px, border-radius 2px. Hover to maroon-deep with an animated arrow (translate 3px on hover). Ghost — transparent, maroon text, 1px line border. Search chip — the one pill in the system."),
    ("Skip link", "Present on every page — a.skip with href #main. Hidden off-screen; on focus appears at top-left on a maroon background. Non-negotiable for accessibility."),
]
for name, body in components:
    add_heading(name, level=2)
    add_para(body, size=11, color=GRAPHITE)

add_heading("Component checklist for a new page", level=2)
add_bullets([
    "Skip link at the top",
    "Utility bar",
    "Masthead with crest and language toggle",
    "Primary navigation with 3px maroon anchor",
    "Section headings with kicker, H2, sub, optional see-more",
    "Buttons only from the primary / ghost / search set",
    "All content inside the 1200px wrap",
    "Alternating paper / paper-2 backgrounds",
    "Footer with maroon-ink field and gold top border",
    "Every image has alt text",
    "Every third-party image has referrerpolicy=no-referrer",
    "No autoplay, no marquee, no infinite scroll",
])

# ==================== 9. MOTION ====================
add_kicker("Chapter Nine")
add_heading("Motion serves reading, never itself.", level=0)
add_para(
    "The site uses motion as punctuation. It is not a stage. Every animated property has a defined role and a defined duration.",
    size=12, color=GRAPHITE, font="Cambria"
)
add_heading("The permitted motions", level=2)
add_bullets([
    "Opacity crossfade — 900ms ease, for hero carousel slides.",
    "Slow image zoom — 10s ease-out, transform scale 1.01 to 1.05, on active slide only.",
    "Hover translate — 3px up for cards, 4px right for directory rows, 180ms ease.",
    "Hover image zoom — scale 1.03 to 1.04, 500ms to 6s ease, on gallery and mosaic cards.",
    "Colour transition — 180ms ease, on links, buttons, and border swaps.",
    "Progress bar — linear, matching slide duration, in gold.",
])
add_heading("Reduced motion", level=2)
add_para(
    "All non-essential motion is suspended under @media (prefers-reduced-motion: reduce). The hero carousel stops zooming; the progress bar is hidden; slide transitions become instantaneous.",
    size=11, color=GRAPHITE
)
add_do_dont(
    [
        "Use the durations above verbatim.",
        "Honour prefers-reduced-motion on every new animation.",
        "Prefer opacity and translate over complex keyframes.",
    ],
    [
        "Never autoplay carousels without arrows, dots, and progress.",
        "Never introduce parallax scroll, mouse-tracking, or sticky video backgrounds.",
        "Never animate text on entry — the reader arrived to read, not to wait.",
        "Never use a marquee ticker. Notices are shown statically.",
    ],
)

# ==================== 10. ACCESSIBILITY ====================
add_kicker("Chapter Ten")
add_heading("WCAG 2.1 AA, held as the floor.", level=0)
add_para(
    "The site is a public institutional face. Its accessibility target is WCAG 2.1 AA — not as an aspiration, as an obligation. Any regression against the checklist below is a bug, not a design decision.",
    size=12, color=GRAPHITE, font="Cambria"
)
add_heading("Structure", level=2)
add_bullets([
    "Every page has one <main id='main'> landmark.",
    "Every page has a skip link (a.skip) pointing at #main.",
    "Landmarks are used correctly — <header>, <nav aria-label>, <main>, <footer>.",
    "Headings are hierarchical. There is exactly one H1 per page.",
])
add_heading("Keyboard", level=2)
add_bullets([
    "Every interactive element is reachable and operable with the keyboard.",
    "Focus rings are visible — the gold outline pattern is the standard.",
    "The carousel arrows and dots are focus-visible; the escape key closes overlays.",
    "No focus trap outside a modal.",
])
add_heading("Content", level=2)
add_bullets([
    "Every image has an alt. Decorative images use alt=''.",
    "Every form field has a real label. Placeholder is not a label.",
    "Every icon-only button carries an aria-label.",
    "Right-click, copy, and text selection are never blocked. Institutional content must be quotable, printable, and screen-readable.",
])
add_heading("Colour and contrast", level=2)
add_bullets([
    "Body text meets 4.5:1. Large text meets 3:1. The pairings in Chapter Four are the verified set.",
    "Colour is never the sole carrier of meaning. Active nav also has a bold weight; error states also have an icon or label.",
])
add_heading("Motion", level=2)
add_bullets([
    "prefers-reduced-motion is honoured on every animation.",
    "Nothing autoplays audio. Nothing flashes more than three times per second.",
])
add_heading("Language", level=2)
add_bullets([
    "The document root sets lang='en' (or lang='hi').",
    "Every Hindi run is marked lang='hi'. The CSS then switches to Devanagari families automatically.",
])
add_para(
    "Accessibility fixes are never a phase two. Ship-blocking, always.",
    size=11, color=MAROON_INK, italic=True, font="Cambria"
)

# ==================== 11. BILINGUAL ====================
add_kicker("Chapter Eleven")
add_heading("Hindi at parity, not decoration.", level=0)
add_para(
    "The site is bilingual by default. Hindi is not a translation appended after the fact; it is present where English is present, and its typography is treated with the same care.",
    size=12, color=GRAPHITE, font="Cambria"
)
add_heading("Where Hindi appears", level=2)
add_bullets([
    "The bilingual wordmark in the masthead.",
    "The bilingual hero H1.",
    "The utility bar language toggle (EN / हिन्दी).",
    "Section-specific runs where the source is Hindi (event names, mottos, motto notes).",
])
add_heading("Marking Hindi in the DOM", level=2)
add_para(
    "Every Hindi run carries lang='hi'. The CSS handles the rest — [lang='hi'] switches to the Devanagari sans; headings with lang='hi' switch to the Devanagari serif. Do not add Hindi text without the lang attribute — screen readers depend on it.",
    size=11, color=GRAPHITE
)
add_heading("Typographic parity", level=2)
add_bullets([
    "Hindi headings are set in Noto Serif Devanagari at a size no more than one step below the Latin heading beside them.",
    "Hindi body is set in IBM Plex Sans Devanagari, at 17px with slightly increased line-height (1.65 vs 1.55 for Latin).",
    "Kickers in Hindi are set in the Devanagari sans, not the Latin.",
])
add_heading("Language toggle", level=2)
add_bullets([
    "Sits in the utility bar, always visible.",
    "Two buttons — EN and हिन्दी. The active button is marked in gold-soft with a 1px gold underline.",
    "Toggling switches the document’s lang attribute and the language of content collections.",
])
add_do_dont(
    [
        "Mark every Hindi run with lang='hi'.",
        "Set the Hindi wordmark at 15px in the masthead (matched to the Latin block).",
        "Increase line-height for Devanagari — the script needs a little more room.",
    ],
    [
        "Never render Hindi in the Latin serif — the shapes are wrong.",
        "Never set Hindi smaller than the corresponding English on the same line.",
        "Never machine-translate institutional content. Hindi content is authored, not generated.",
    ],
)

# ==================== 12. VOICE ====================
add_kicker("Chapter Twelve")
add_heading("The voice of the site.", level=0)
add_para(
    "The site’s writing style is Indian formal English. It is grammatical, plain, and confident. It uses the third-person institutional voice by default. It does not adopt the register of a startup, a marketing brochure, or a private consultancy.",
    size=12, color=GRAPHITE, font="Cambria"
)
add_heading("Spelling and register", level=2)
add_bullets([
    "British spellings — organise, recognise, programme, honour, centre.",
    "No colloquialisms — never buy-in, low-hanging fruit, deep dive, circle back, reach out, ping, leverage (as a verb).",
    "No unearned tech jargon — no stack, platform (as buzzword), ecosystem, primitive.",
    "No dramatic bold, no exclamation marks, no all-caps headlines.",
])
add_heading("Names, dates, numbers", level=2)
add_bullets([
    "The institution is Govind Ballabh Pant University of Agriculture and Technology (formal), or GBPUA&T, or Pantnagar. Do not shorten to ‘the university’ without antecedent.",
    "Dates in institutional prose are spelled — ‘13–16 March 2026’ — with an en dash for ranges.",
    "Numerals: figures for exact quantities and years; words for one to nine in prose.",
    "Punctuation is smart. Use — for em dash, – for en dash, ’ for the apostrophe. Straight quotes are a bug.",
])
add_heading("Content is preserved", level=2)
add_para(
    "Article text, dates, and figures pulled from the live site or provided by the institution are used verbatim. The website team does not paraphrase institutional content. If a source is unclear, ask the content owner — do not smooth it.",
    size=11, color=GRAPHITE
)
add_heading("Voice for microcopy", level=2)
add_bullets([
    "CTAs are noun phrases or short verb phrases — ‘Admissions 2026–27’, ‘Read more’, ‘See all notices’.",
    "Kickers are institutional or temporal — ‘Since 1960’, ‘Notices’, ‘This week’.",
    "Empty states are literal — ‘No circulars this week.’ Not clever.",
])

# ==================== 13. GOVERNANCE ====================
add_kicker("Chapter Thirteen")
add_heading("How the system stays whole.", level=0)
add_para(
    "A design system is not defended by a document. It is defended by a small set of habits held by the team maintaining it.",
    size=12, color=GRAPHITE, font="Cambria"
)
add_heading("Tokens are edited in one place", level=2)
add_para(
    "All design tokens live in the :root block at the top of the site’s stylesheet. Colour, font, spacing, rules — nothing is hand-coded elsewhere. If a new colour or size is needed, it is added to the tokens, then referenced. This is the single most important rule in the document.",
    size=11, color=GRAPHITE
)
add_heading("Components are reused before invented", level=2)
add_para(
    "Before writing a new component, walk Chapter Eight and ask whether an existing component can be adapted. A new component is added only when nothing fits. When one is added, it is added to Chapter Eight in the next revision.",
    size=11, color=GRAPHITE
)
add_heading("New pages copy an existing page’s skeleton", level=2)
add_para(
    "The homepage is the reference. Inner pages start from it and simplify. This keeps the design language consistent as the site grows across colleges, directorates, and cells.",
    size=11, color=GRAPHITE
)
add_heading("Every change is checked against the checklist", level=2)
add_para(
    "Chapter Eight closes with a checklist for a new page. Chapter Ten sets the accessibility floor. Chapters Four and Five set the token constraints. Before a change goes live, it is walked against these three lists.",
    size=11, color=GRAPHITE
)
add_heading("The AS-IS Institutional Brand Guideline governs the identity", level=2)
add_para(
    "This document governs the digital expression. If a question arises about the crest, the motto, the formal use of the wordmark, or the identity in any medium other than the website, the AS-IS guideline is authoritative. When in doubt, ask, and record the answer here.",
    size=11, color=GRAPHITE
)
add_heading("The document is revised, not rewritten", level=2)
add_para(
    "New components, new pages, and new lessons are added as they arrive. Old sections are corrected when the site departs from them. Version numbers are bumped when the palette or the type families change; minor revisions do not need a new number.",
    size=11, color=GRAPHITE
)
add_para(
    "The site is a living system. This document is the record of that system. The record must stay honest — if the site departs from it, the record is updated, or the site is corrected. Divergence is a bug.",
    size=11, color=MAROON_INK, italic=True, font="Cambria"
)

# Save
import os
from pathlib import Path
out = str(Path(__file__).parent / "DIGITAL-BRAND-GUIDELINE.docx")
doc.save(out)
print(f"Wrote {out} ({os.path.getsize(out)/1024:.1f} KB)")
